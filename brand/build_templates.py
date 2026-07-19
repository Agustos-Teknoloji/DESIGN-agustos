#!/usr/bin/env python3
"""
Ağustos brand kit — office documents, swatches, email signature, guidelines source.

Sibling to build.py (which makes the visual logo assets). This generates the
"working document" deliverables from the same registry, so they stay in sync:

  exports/<brand>/swatches/    <brand>.ase (Adobe) + <brand>.clr (Apple)
  exports/<brand>/email/       <brand>-signature.html (email-safe, self-contained)
  exports/<brand>/office/      <brand>-letterhead.docx + <brand>-document-template.docx + <brand>-template.pptx
  exports/<brand>/guidelines/  <brand>-brand-guidelines.html (rendered to PDF by browse)

Run after build.py (it reuses the generated lockup PNGs).

  ../.venv/bin/python build_templates.py [--brand <slug>]
"""

from __future__ import annotations

import argparse
import base64
import io
import json
import subprocess
import struct
import sys
from pathlib import Path

from PIL import Image

BRAND_DIR = Path(__file__).resolve().parent
ROOT = BRAND_DIR.parent
REGISTRY = BRAND_DIR / "brands.json"
TOKENS = ROOT / "tokens" / "resolved.json"


def hexrgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i + 2], 16) for i in (0, 2, 4))


def b64_png(path: Path, width: int) -> str:
    """Resize a PNG to `width` and return a base64 data URI (for self-contained email)."""
    im = Image.open(path).convert("RGBA")
    h = round(im.height * width / im.width)
    im = im.resize((width, h), Image.LANCZOS)
    buf = io.BytesIO()
    im.save(buf, format="PNG")
    return "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode()


# ----------------------------------------------------------------------------
# Color swatches
# ----------------------------------------------------------------------------

def palette(brand, reg, design):
    """Identity ink, shared interaction signal, and system neutrals."""
    colors = design["foundations"]["color"]
    return [
        (f'{brand["title"]} Identity', brand["color"]),
        ("Shared Signal Red", design["semantic"]["color"]["signal"]),
        ("Ink", colors["ink"]),
        ("Ink Soft", colors["inkSoft"]),
        ("Ink Faint", colors["inkFaint"]),
        ("Paper (Cream)", colors["paperCream"]),
        ("Paper White", colors["paperWhite"]),
        ("Rule", colors["ruleCream"]),
    ]


def write_ase(colors, out: Path):
    """Adobe Swatch Exchange (.ase) — opens in Illustrator, Photoshop, InDesign,
    Affinity, Figma (via plugin). Big-endian binary."""
    blob = b"ASEF" + struct.pack(">HH", 1, 0) + struct.pack(">I", len(colors))
    for name, hx in colors:
        r, g, b = (c / 255 for c in hexrgb(hx))
        nm = name.encode("utf-16-be") + b"\x00\x00"
        body = struct.pack(">H", len(name) + 1) + nm + b"RGB " + struct.pack(">fff", r, g, b) + struct.pack(">H", 2)
        blob += struct.pack(">H", 1) + struct.pack(">I", len(body)) + body
    out.write_bytes(blob)


def write_clr(colors, list_name, out: Path):
    """Apple Color List (.clr) — drops into the macOS system color picker.
    Uses AppKit (PyObjC); macOS only."""
    from AppKit import NSColorList, NSColor
    from Foundation import NSURL
    cl = NSColorList.alloc().initWithName_(list_name)
    for name, hx in colors:
        r, g, b = (c / 255 for c in hexrgb(hx))
        cl.setColor_forKey_(NSColor.colorWithSRGBRed_green_blue_alpha_(r, g, b, 1.0), name)
    cl.writeToURL_error_(NSURL.fileURLWithPath_(str(out)), None)


# ----------------------------------------------------------------------------
# Email signature (email-safe HTML: table layout, inline styles, web-safe font)
# ----------------------------------------------------------------------------

def gen_email_signature(slug, brand, reg, design, out: Path, lockup_png: Path):
    logo = b64_png(lockup_png, 320)
    signal = design["semantic"]["color"]["signal"]
    title = brand["title"]
    domain = brand.get("domain", "")
    # Tagline intentionally omitted — defined in the registry but used sparingly,
    # never on everyday artifacts (see brands.json $tagline_policy).
    html = f"""<!-- {title} email signature. Paste into your mail client's signature editor.
     Self-contained (logo embedded). Replace {{{{NAME}}}}, {{{{ROLE}}}}, {{{{PHONE}}}}. -->
<table cellpadding="0" cellspacing="0" border="0" style="font-family:Arial,Helvetica,sans-serif;color:#1a1a1a;">
  <tr>
    <td style="padding-right:18px;vertical-align:middle;">
      <img src="{logo}" width="150" alt="{title}" style="display:block;border:0;">
    </td>
    <td style="border-left:2px solid {signal};padding-left:18px;vertical-align:middle;line-height:1.5;">
      <div style="font-size:15px;font-weight:bold;color:#1a1a1a;">{{{{NAME}}}}</div>
      <div style="font-size:13px;color:#4a4a4a;padding-bottom:6px;">{{{{ROLE}}}} &middot; {title}</div>
      <div style="font-size:12px;color:#4a4a4a;">
        {{{{PHONE}}}} &nbsp;|&nbsp;
        <a href="mailto:hello@{domain}" style="color:#4a4a4a;text-decoration:none;">hello@{domain}</a> &nbsp;|&nbsp;
        <a href="https://{domain}" style="color:{signal};text-decoration:none;font-weight:bold;">{domain}</a>
      </div>
    </td>
  </tr>
</table>
"""
    out.write_text(html, encoding="utf-8")


# ----------------------------------------------------------------------------
# Word letterhead (.docx)
# ----------------------------------------------------------------------------

def _remove_paragraph_border(element):
    """Keep Word/Google Docs title imports free of built-in rule residue."""
    from docx.oxml.ns import qn
    p_pr = element.find(qn("w:pPr"))
    if p_pr is None:
        return
    for border in list(p_pr.findall(qn("w:pBdr"))):
        p_pr.remove(border)


def _measure(value: str, unit: str) -> float:
    if not value.endswith(unit):
        raise ValueError(f"expected {unit} value, got {value!r}")
    return float(value[:-len(unit)])


def _set_cell_shading(cell, color: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color.lstrip("#"))
    tc_pr.append(shading)


def _set_cell_bottom_border(cell, color: str, size="12"):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    for key, value in (("w:val", "single"), ("w:sz", size), ("w:color", color.lstrip("#"))):
        bottom.set(qn(key), value)
    borders.append(bottom)
    tc_pr.append(borders)


def gen_document_docx(slug, brand, design, out: Path, lockup_png: Path, *, letterhead=False):
    import docx
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

    foundations = design["foundations"]
    recipe = design["recipes"]["document"]
    colors = foundations["color"]
    display = foundations["fontFamily"]["display"]
    body = foundations["fontFamily"]["body"]
    ink = RGBColor(*hexrgb(colors["ink"]))
    ink_soft = RGBColor(*hexrgb(colors["inkSoft"]))
    signal_hex = design["semantic"]["color"]["signal"]
    signal_color = RGBColor(*hexrgb(signal_hex))
    title, domain = brand["title"], brand.get("domain", "")

    doc = docx.Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    margin_cm = _measure(recipe["margin"], "mm") / 10
    for m in ("bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(margin_cm))
    sec.top_margin = Cm(_measure(recipe["topMargin"], "mm") / 10)

    normal = doc.styles["Normal"]
    normal.font.name = body
    normal.font.size = Pt(_measure(recipe["bodySize"], "pt"))
    normal.font.color.rgb = ink
    normal.paragraph_format.space_after = Pt(_measure(recipe["bodyAfter"], "pt"))
    normal.paragraph_format.line_spacing = recipe["bodyLineSpacing"]

    style_specs = {
        "Title": (recipe["titleSize"], 600, ink, 12),
        "Heading 1": (recipe["heading1Size"], 600, ink, 9),
        "Heading 2": (recipe["heading2Size"], 600, ink, 7),
        "Heading 3": (recipe["heading3Size"], 650, signal_color, 5),
    }
    for name, (size, weight, rgb, after) in style_specs.items():
        style = doc.styles[name]
        if name == "Title":
            _remove_paragraph_border(style.element)
        style.font.name = display
        style.font.size = Pt(_measure(size, "pt"))
        style.font.bold = weight >= 600
        style.font.color.rgb = rgb
        style.paragraph_format.space_before = Pt(12 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Ağustos Label" not in doc.styles:
        label = doc.styles.add_style("Ağustos Label", WD_STYLE_TYPE.PARAGRAPH)
    else:
        label = doc.styles["Ağustos Label"]
    label.font.name = display
    label.font.size = Pt(_measure(recipe["footnoteSize"], "pt"))
    label.font.bold = True
    label.font.color.rgb = signal_color
    label.paragraph_format.space_after = Pt(5)

    hdr = sec.header
    hp = hdr.paragraphs[0]
    hp.add_run().add_picture(str(lockup_png), width=Cm(4.6))
    rule = hdr.add_table(rows=1, cols=1, width=Cm(16.6))
    rule_cell = rule.cell(0, 0)
    _set_cell_bottom_border(rule_cell, signal_hex)
    rule_p = rule_cell.paragraphs[0]
    rule_p.paragraph_format.space_before = Pt(0)
    rule_p.paragraph_format.space_after = Pt(0)
    rule_p.add_run(" ").font.size = Pt(1)

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(f"{title}  ·  {domain}")
    r.font.size = Pt(8.5)
    r.font.color.rgb = ink_soft

    if letterhead:
        doc.add_paragraph("Letter title", style="Title")
        doc.add_paragraph(
            "Replace this with your letter body. The header, signal rule, and footer are part "
            "of the letterhead. Use the built-in styles so the document keeps the shared rhythm."
        )
    else:
        doc.add_paragraph("DOCUMENT TEMPLATE", style="Ağustos Label")
        doc.add_paragraph("A clear opening for substantial work.", style="Title")
        deck = doc.add_paragraph("Subtitle, author, or date")
        deck.style = doc.styles["Subtitle"]
        deck.runs[0].font.name = body
        deck.runs[0].font.color.rgb = ink_soft
        doc.add_paragraph(
            "This file carries the same alignment, hierarchy, signal color, and spacing logic as "
            "the website without copying a web page into a document. It imports cleanly into Google Docs."
        )
        doc.add_paragraph("Section opening", style="Heading 1")
        doc.add_paragraph(
            "Use Heading 1 to begin a major section. Body copy stays calm and readable; shared red "
            "is reserved for links, labels, rules, and decisive signals."
        )
        doc.add_paragraph("A smaller hierarchy", style="Heading 2")
        doc.add_paragraph("Use Heading 2 and Heading 3 to create structure without adding decoration.")
        doc.add_paragraph("Structured information", style="Heading 3")
        table = doc.add_table(rows=3, cols=3)
        table.style = "Table Grid"
        headers = ("Role", "System value", "Use")
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_shading(cell, brand["color"])
            for run in cell.paragraphs[0].runs:
                run.font.name = display
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFE, 0xFC, 0xF2)
        rows = (
            ("Identity ink", brand["color"], "Logo and decisive brand fields"),
            ("Shared signal", signal_hex, "Links, focus, markers, and emphasis"),
        )
        for row_index, values in enumerate(rows, start=1):
            for col_index, text in enumerate(values):
                table.rows[row_index].cells[col_index].text = text
    doc.save(str(out))


# ----------------------------------------------------------------------------
# PowerPoint template (.pptx) — native editable objects also import to Google Slides
# ----------------------------------------------------------------------------

def gen_pptx(slug, out: Path):
    command = ["node", str(BRAND_DIR / "build_presentation.mjs"), "--brand", slug, "--output", str(out)]
    try:
        subprocess.run(command, check=True)
    except subprocess.CalledProcessError as error:
        raise SystemExit(
            "PowerPoint generation requires the declared brand Node dependencies. "
            "Run npm install in brand/ and retry."
        ) from error


# ----------------------------------------------------------------------------
# Brand guidelines (HTML — rendered to PDF by browse)
# ----------------------------------------------------------------------------

def gen_guidelines_html(slug, brand, reg, design, out: Path, lk_dir: Path, fav_dir: Path):
    color = brand["color"]
    signal = design["semantic"]["color"]["signal"]
    title, domain = brand["title"], brand.get("domain", "")
    fonts = BRAND_DIR / "fonts"
    pal = palette(brand, reg, design)
    sw = "".join(
        f'<div class="sw"><div class="chip" style="background:{hx};'
        f'{"border:1px solid #e8e3d0;" if hx.lower() in ("#fefcf2","#ffffff") else ""}"></div>'
        f'<div class="swn">{name}</div><div class="swh">{hx.upper()}</div></div>'
        for name, hx in pal
    )
    # file:// URLs to the bundled fonts so the PDF renders in real brand type.
    it = (fonts / "inter-tight" / "InterTight[wght].ttf").as_uri()
    inr = (fonts / "inter" / "Inter[opsz,wght].ttf").as_uri()
    pos = (lk_dir / f"{slug}-lockup__positive.svg").as_uri()
    neg = (lk_dir / f"{slug}-lockup__negative.svg").as_uri()
    mono = (lk_dir / f"{slug}-lockup__mono.svg").as_uri()
    fav = (fav_dir / "favicon.svg").as_uri()

    html = f"""<!doctype html><html><head><meta charset="utf-8"><style>
@font-face {{ font-family:'IT'; src:url('{it}'); }}
@font-face {{ font-family:'IN'; src:url('{inr}'); }}
@page {{ size:A4; margin:0; }}
* {{ box-sizing:border-box; -webkit-print-color-adjust:exact; print-color-adjust:exact; }}
body {{ margin:0; font-family:'IN',sans-serif; color:#1a1a1a; background:#fefcf2; }}
.page {{ width:210mm; min-height:297mm; padding:22mm 20mm; page-break-after:always; position:relative; }}
.page:last-child {{ page-break-after:auto; }}
h1 {{ font-family:'IT'; font-weight:650; font-size:46px; letter-spacing:-0.03em; margin:0 0 6px; }}
h2 {{ font-family:'IT'; font-weight:650; font-size:13px; text-transform:uppercase; letter-spacing:0.14em; color:#8a8a8a; margin:34px 0 14px; }}
p {{ font-size:13.5px; line-height:1.65; max-width:62ch; color:#4a4a4a; }}
.cover-mark {{ width:230px; margin:48mm 0 10mm; }}
.eyebrow {{ font-family:'IT'; font-weight:650; font-size:12px; text-transform:uppercase; letter-spacing:0.16em; color:{signal}; }}
.foot {{ position:absolute; bottom:14mm; left:20mm; right:20mm; font-size:10px; color:#8a8a8a; border-top:1px solid #e8e3d0; padding-top:6px; display:flex; justify-content:space-between; }}
.row {{ display:flex; gap:18px; flex-wrap:wrap; align-items:flex-end; }}
.card {{ border:1px solid #e8e3d0; border-radius:6px; padding:18px; }}
.card.dark {{ background:{color}; border-color:{color}; }}
.card img {{ height:38px; display:block; }}
.lbl {{ font-family:'IT'; font-weight:650; font-size:10px; text-transform:uppercase; letter-spacing:0.1em; color:#8a8a8a; margin-top:12px; }}
.swatches {{ display:grid; grid-template-columns:repeat(4,1fr); gap:14px; }}
.chip {{ height:64px; border-radius:6px; }}
.swn {{ font-family:'IT'; font-weight:650; font-size:12px; margin-top:8px; }}
.swh {{ font-size:11px; color:#8a8a8a; font-variant-numeric:tabular-nums; }}
.type-it {{ font-family:'IT'; }}
.spec {{ font-size:32px; }}
.do {{ color:#1f6b4a; font-weight:bold; }} .dont {{ color:#b42318; font-weight:bold; }}
ul.rules {{ font-size:12.5px; line-height:1.7; color:#4a4a4a; padding-left:18px; }}
.clearbox {{ display:inline-block; border:1px dashed {signal}; padding:14px; }}
.clearbox img {{ height:46px; display:block; }}
</style></head><body>

<section class="page">
  <div class="eyebrow">Brand Guidelines</div>
  <img class="cover-mark" src="{pos}">
  <h1>{title}</h1>
  <div class="foot"><span>{domain}</span><span>v1 · generated from the Ağustos Design System</span></div>
</section>

<section class="page">
  <div class="eyebrow">01</div><h1>The logo</h1>
  <p>The mark is the Laz Güneşi symbol locked up with the wordmark. Three expressions cover every
     surface. Use the positive form first — it works on cream and dark backgrounds alike.</p>
  <div class="row" style="margin-top:18px;">
    <div class="card"><img src="{pos}"><div class="lbl">Positive — primary</div></div>
    <div class="card dark"><img src="{neg}"><div class="lbl" style="color:rgba(255,255,255,.7)">Negative — on brand</div></div>
    <div class="card"><img src="{mono}"><div class="lbl">Mono — single colour</div></div>
  </div>
  <h2>Clear space &amp; minimum size</h2>
  <p>Keep clear space around the lockup equal to the symbol height. Never place type or edges inside it.
     Minimum size: 24px tall on screen, 8mm in print, so the wordmark stays legible.</p>
  <div class="clearbox"><img src="{pos}"></div>
  <div class="foot"><span>{title} — Brand Guidelines</span><span>{domain}</span></div>
</section>

<section class="page">
  <div class="eyebrow">02</div><h1>Colour</h1>
  <p>Identity ink names the brand: Ağustos is red; every other house brand is black and white.
     Shared Ağustos red is the interaction signal for links, focus, markers, rules, and small accents.</p>
  <div class="swatches" style="margin-top:18px;">{sw}</div>
  <h2>Do &amp; don't</h2>
  <ul class="rules">
    <li><span class="do">Do</span> set the wordmark in Inter Tight Semibold, lowercase, in the registered identity ink.</li>
    <li><span class="dont">Don't</span> recolour the symbol, stretch the lockup, or add a tagline beside it.</li>
    <li><span class="do">Do</span> use shared red for interaction and emphasis; <span class="dont">don't</span> recolour a non-Ağustos logo red.</li>
  </ul>
  <div class="foot"><span>{title} — Brand Guidelines</span><span>{domain}</span></div>
</section>

<section class="page">
  <div class="eyebrow">03</div><h1>Typography</h1>
  <p>Two families, one skeleton. Inter Tight for display and the wordmark; Inter for body. JetBrains Mono for code and data.</p>
  <div style="margin-top:22px;">
    <div class="type-it spec" style="font-weight:650;">Inter Tight Semibold</div>
    <div class="lbl">Display · wordmark · headings</div>
  </div>
  <div style="margin-top:26px;">
    <div class="spec">Inter — body text</div>
    <div class="lbl">Paragraphs · captions · tables</div>
  </div>
  <p style="margin-top:30px;">The Turkish locale is first-class: <span class="type-it" style="font-weight:650;">ağustos</span>,
     <span class="type-it" style="font-weight:650;">İstanbul</span>, <span class="type-it" style="font-weight:650;">ışık</span> —
     the dotted/dotless i and the ğ breve all render correctly.</p>
  <div class="foot"><span>{title} — Brand Guidelines</span><span>{domain}</span></div>
</section>

</body></html>"""
    out.write_text(html, encoding="utf-8")


# ----------------------------------------------------------------------------

def build_brand(slug, brand, reg, design):
    base = BRAND_DIR / "exports" / slug
    lk = base / "lockup"
    fav = base / "favicon"
    pos_png = lk / f"{slug}-lockup__positive.png"
    if not pos_png.exists():
        raise SystemExit(f"missing {pos_png} — run build.py --brand {slug} first")

    sw_dir = base / "swatches"; sw_dir.mkdir(parents=True, exist_ok=True)
    em_dir = base / "email"; em_dir.mkdir(parents=True, exist_ok=True)
    of_dir = base / "office"; of_dir.mkdir(parents=True, exist_ok=True)
    gl_dir = base / "guidelines"; gl_dir.mkdir(parents=True, exist_ok=True)

    pal = palette(brand, reg, design)
    write_ase(pal, sw_dir / f"{slug}.ase")
    try:
        write_clr(pal, brand["title"], sw_dir / f"{slug}.clr")
    except Exception as e:
        print(f"    (skipped .clr: {e})")
    gen_email_signature(slug, brand, reg, design, em_dir / f"{slug}-signature.html", pos_png)
    gen_document_docx(slug, brand, design, of_dir / f"{slug}-letterhead.docx", pos_png, letterhead=True)
    gen_document_docx(slug, brand, design, of_dir / f"{slug}-document-template.docx", pos_png)
    gen_pptx(slug, of_dir / f"{slug}-template.pptx")
    gen_guidelines_html(slug, brand, reg, design, gl_dir / f"{slug}-brand-guidelines.html", lk, fav)

    n = sum(1 for _ in base.rglob("*") if _.is_file())
    print(f"  ✓ {slug}: office + extras done ({n} files total in exports/{slug})")


def main():
    reg = json.loads(REGISTRY.read_text(encoding="utf-8"))
    if not TOKENS.exists():
        raise SystemExit("missing tokens/resolved.json — run scripts/build_design_system.py first")
    design = json.loads(TOKENS.read_text(encoding="utf-8"))
    ap = argparse.ArgumentParser()
    ap.add_argument("--brand")
    args = ap.parse_args()
    brands = reg["brands"]
    if args.brand and args.brand not in brands:
        raise SystemExit(f"unknown brand '{args.brand}'")
    if args.brand:
        if args.brand not in brands:
            raise SystemExit(f"unknown brand: {args.brand}")
        if not brands[args.brand].get("office", False):
            raise SystemExit(f"{args.brand} does not have an Office kit")
        targets = {args.brand: brands[args.brand]}
    else:
        targets = {slug: brand for slug, brand in brands.items() if brand.get("office", False)}
    print(f"Building office/extras for {len(targets)} brand(s)...")
    for slug, brand in targets.items():
        build_brand(slug, brand, reg, design)
    if args.brand:
        print("  Office manifest unchanged; run a full build before release")
    else:
        subprocess.run(
            [sys.executable, str(ROOT / "scripts" / "check_office_artifacts.py"), "--write"],
            cwd=ROOT,
            check=True,
        )
    print("Done. (Render guidelines HTML -> PDF with browse; see templates/README.md)")


if __name__ == "__main__":
    main()
